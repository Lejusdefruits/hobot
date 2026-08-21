"""Per-offer CV tailoring: edits the candidate's own uploaded CV in place --
same file, same layout, same fonts, same colors -- rather than regenerating
a document from a template. Only two things ever change: the profile/summary
paragraph, and (if the CV's skills section is a plain text list, not an
icon/pill layout) which of the candidate's real skills fill the visible
slots. Nothing else -- name, contact info, dates, job titles, company names,
degree names, images, colors, section order -- is ever touched. Never
invents a skill, a project, or an experience not already in user_profile.

Dispatches on user_profile.cv_format:
- pdf_text: the CV has a real, extractable text layer. Redact the target
  text region and reinsert the new text with matched font/size/color -- the
  common, highest-fidelity case.
- docx: the CV is a Word document. Renders through a Jinja2 tag inserted
  into the source once (see _ensure_docx_template), then docxtpl fills it
  per offer and LibreOffice converts the result to PDF.
- pdf_image: the CV is a flattened raster export (common for Canva "PDF for
  print" downloads) with zero extractable text. True in-place editing is
  impossible here -- there's no text object to redact. The realistic,
  honest version of this feature for such a CV is: locate the summary
  paragraph's region with one narrowly-scoped vision call, cover it, and
  insert the new text as a REAL PDF text object (not repainted pixels) --
  the point of tailoring is ATS keyword matching, and an ATS parser gets
  nothing from pixels. This branch cannot promise the same visual fidelity
  as the other two and says so in its own return value.
"""
import os
import re
import subprocess
import tempfile
import uuid
from collections import Counter
from pathlib import Path

import pymupdf as fitz

from core.db import get_connection
from core.llm import chat_json
from core.profile import PROFILE_DIR
from tools.common import normalize_text

OUTPUT_DIR = Path(__file__).resolve().parent.parent / os.environ.get("HOBOT_OUTPUT_DIR", "outputs")
FONTS_DIR = Path(__file__).resolve().parent.parent / "assets" / "fonts"

# A rewritten summary paragraph must stay close to the original length -- a
# CV is laid out for a fixed page, a longer paragraph pushes everything else
# out of position. LLMs unreliably ignore a length instruction in the prompt
# alone (confirmed in practice, not theoretical), so this is also enforced
# programmatically as a hard fallback, never prompt-only.
SUMMARY_MAX_CHARS = int(os.environ.get("CV_TAILOR_SUMMARY_MAX_CHARS", "500"))

SUMMARY_TAILOR_PROMPT = """You are rewriting the "profile" or "summary" paragraph at the top of a
candidate's CV for ONE specific job posting, without ever inventing anything.

Current summary (never stray from the facts it contains):
{original}

Candidate's real skills: {skills}
Candidate's real experience: {experience}
Candidate's real education: {education}

Target posting:
- Title: {title}
- Description: {description}

Rewrite this paragraph (same tone, same language as the current summary) highlighting whichever
REAL elements above best fit this posting, echoing the posting's own wording where that's honest
to do. NEVER invent a skill, project, experience, or degree that isn't listed above -- a
paragraph that stays close to the original beats an inaccurate one.

STRICT LIMIT: {max_chars} characters maximum, spaces included -- the original is {original_chars}
characters, do not exceed it. If you have to choose, a shorter, cleaner paragraph beats a
complete but too-long one.

Reply with ONLY JSON: {{"summary": "..."}}"""

SKILL_PICK_PROMPT = """A candidate's CV lists these skills, some visible on the CV page, some only
known from their fuller profile (added later, not yet on the visible page):

Visible on the CV right now: {visible}
Known but not currently visible: {hidden}

Target posting:
- Title: {title}
- Description: {description}

Pick UP TO {max_picks} skill(s) from the "known but not currently visible" list that are
genuinely relevant to this posting and worth surfacing -- never invent a new one, never pick
one that isn't in that list verbatim. If none of them fit, return an empty list.

Reply with ONLY JSON: {{"skills": ["..."]}}"""


def _full_profile(conn) -> dict:
    import json
    row = conn.execute(
        "SELECT skills, experience, education, cv_source_path, cv_format FROM user_profile WHERE id = 1"
    ).fetchone()
    if not row:
        return {}
    return {
        "skills": json.loads(row["skills"] or "[]"),
        "experience": json.loads(row["experience"] or "[]"),
        "education": json.loads(row["education"] or "[]"),
        "cv_source_path": row["cv_source_path"],
        "cv_format": row["cv_format"],
    }


def _truncate_to_sentence(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    cut = text[:max_chars]
    end = max(cut.rfind(". "), cut.rfind("? "), cut.rfind("! "))
    return cut[: end + 1] if end > max_chars * 0.5 else cut.rstrip() + "…"


def _tailor_summary_text(original: str, profile: dict, title: str, description: str) -> str:
    """One LLM call, budget enforced both in the prompt and programmatically
    afterward -- never trust the prompt instruction alone. Falls back to the
    untouched original on any failure or empty result, same as this project's
    other best-effort generation steps (a failed rewrite must never block the
    rest of the pipeline)."""
    try:
        result = chat_json(SUMMARY_TAILOR_PROMPT.format(
            original=original,
            max_chars=SUMMARY_MAX_CHARS,
            original_chars=len(original),
            skills=", ".join(profile.get("skills", [])) or "(none)",
            experience="; ".join(
                f"{e.get('poste')} at {e.get('entreprise')} ({e.get('periode')}): {e.get('resume')}"
                for e in profile.get("experience", [])
            ) or "(none)",
            education="; ".join(
                f"{e.get('diplome')} -- {e.get('etablissement')} ({e.get('periode')})"
                for e in profile.get("education", [])
            ) or "(none)",
            title=title or "", description=(description or "")[:2000],
        ))
        summary = (result.get("summary") or "").strip()
        return _truncate_to_sentence(summary, SUMMARY_MAX_CHARS) if summary else original
    except Exception:
        return original


def _pick_hidden_skills(visible: list[str], profile_skills: list[str], title: str, description: str,
                         max_picks: int = 2) -> list[str]:
    """Skills already in user_profile.skills (e.g. added later through chat
    via modifier_profil) that aren't on the CV's visible list yet -- the
    actual valuable move here, since reordering what's already visible has
    close to no real ATS effect (text extraction doesn't weight by visual
    position). Never returns anything not verbatim in `profile_skills`."""
    visible_norm = {normalize_text(v) for v in visible}
    hidden = [s for s in profile_skills if normalize_text(s) not in visible_norm]
    if not hidden:
        return []
    try:
        result = chat_json(SKILL_PICK_PROMPT.format(
            visible=", ".join(visible) or "(none)", hidden=", ".join(hidden),
            title=title or "", description=(description or "")[:1000], max_picks=max_picks,
        ))
        picked = result.get("skills") or []
        return [s for s in picked if s in hidden][:max_picks]
    except Exception:
        return []


# ---------------------------------------------------------------------------
# pdf_text branch
# ---------------------------------------------------------------------------

_SIMPLE_FONTS = {
    "helv": FONTS_DIR / "Montserrat-Regular.ttf",
    "hebo": FONTS_DIR / "Montserrat-Bold.ttf",
}


def _text_blocks(page) -> list[dict]:
    """One entry per visual block, its lines joined in reading order --
    matching a multi-line paragraph against PyMuPDF's raw per-line spans
    almost never works, a paragraph is essentially never one span's exact
    `.text`."""
    blocks = []
    for b in page.get_text("dict")["blocks"]:
        if "lines" not in b:
            continue
        lines = [" ".join(s["text"] for s in l["spans"]) for l in b["lines"]]
        spans = [s for l in b["lines"] for s in l["spans"]]
        if spans:
            blocks.append({"bbox": b["bbox"], "text": " ".join(lines), "spans": spans})
    return blocks


def _font_glyphs_ok(fontfile: Path, text: str) -> bool:
    from fontTools.ttLib import TTFont
    cmap = TTFont(str(fontfile)).getBestCmap()
    return all(ord(c) in cmap for c in text)


def _pick_font(page, spans: list[dict], text: str) -> tuple[Path, float, tuple]:
    """Real embedded font first, checked glyph-by-glyph against the actual
    replacement text -- confirmed in practice (fontTools cmap check against a
    real CV's embedded, subsetted font) that this fails routinely, not
    rarely: a Canva/Word/Docs PDF export typically only embeds the glyphs its
    original content used, so a common letter the original text never needed
    (a 'w', a 'z') is often simply missing. On any miss, fall back to the
    bundled real Montserrat static instance (matches this project's most
    common real-world case) rather than silently guessing a base-14 font.

    Deliberately only DECIDES which font file to use here and does not
    register it on the page yet (no page.insert_font call) -- that has to
    happen after redaction, not before: apply_redactions() rebuilds the
    page's whole content stream, and a font registered before that call
    stops working (confirmed in practice, not theoretical)."""
    ref = spans[0]
    size = ref["size"]
    color_int = ref["color"]
    color = ((color_int >> 16 & 255) / 255, (color_int >> 8 & 255) / 255, (color_int & 255) / 255)

    # extract_font needs an xref, not a name -- PyMuPDF only exposes the font
    # name on the span, so go through get_fonts() to resolve it.
    try:
        for f in page.get_fonts():
            if f[3] == ref["font"]:
                fontbuffer = page.parent.extract_font(f[0])[-1]
                tmp = Path(tempfile.mkdtemp()) / "embedded.ttf"
                tmp.write_bytes(fontbuffer)
                if _font_glyphs_ok(tmp, text):
                    return tmp, size, color
                break
    except Exception:
        pass

    # bundled fallback -- always covers the full glyph set we ship it for
    bold = "bold" in ref["font"].lower() or "semibold" in ref["font"].lower()
    fallback = FONTS_DIR / ("Montserrat-Bold.ttf" if bold else "Montserrat-Regular.ttf")
    return fallback, size, color


def _sample_background(page, bbox: fitz.Rect) -> tuple[int, int, int]:
    """Several sample points at a few different offsets around the block,
    keep the MOST COMMON color rather than an average -- an average gets
    dragged off by a single point that lands on a neighboring element (a
    colored badge's edge sitting right above a paragraph, confirmed on a
    real CV to produce a visibly wrong fill otherwise)."""
    votes = []
    for dy in (-10, -8, -6, bbox.height + 6, bbox.height + 8, bbox.height + 10):
        y = bbox.y0 + dy
        if y < 0 or y > page.rect.height:
            continue
        for dx in (0, bbox.width * 0.3, bbox.width * 0.6, bbox.width * 0.9):
            x = bbox.x0 + dx
            if x + 1 > page.rect.width:
                continue
            pix = page.get_pixmap(clip=fitz.Rect(x, y, x + 1, y + 1))
            if pix.samples:
                votes.append((pix.samples[0], pix.samples[1], pix.samples[2]))
    return Counter(votes).most_common(1)[0][0] if votes else (255, 255, 255)


def _wrap_and_insert(page, bbox: fitz.Rect, text: str, fontfile: Path, fontsize: float, color: tuple) -> None:
    """Registers the font and inserts the text -- must run AFTER any
    redaction on this page (see _pick_font's docstring)."""
    font = fitz.Font(fontfile=str(fontfile))
    fontname = f"cvtailor-{uuid.uuid4().hex[:8]}"
    page.insert_font(fontname=fontname, fontfile=str(fontfile))

    # measuring against the exact loaded font object is what actually
    # matters here -- page.get_text_length() only has accurate metrics for
    # the base-14/CJK built-ins, not an arbitrary embedded or bundled font.
    words = text.split(" ")
    lines, current = [], ""
    for w in words:
        trial = (current + " " + w).strip()
        if font.text_length(trial, fontsize) <= bbox.width:
            current = trial
        else:
            lines.append(current)
            current = w
    if current:
        lines.append(current)
    line_height = fontsize * 1.3
    y = bbox.y0 + fontsize * 0.9
    for line in lines:
        page.insert_text((bbox.x0, y), line, fontname=fontname, fontsize=fontsize, color=color)
        y += line_height


def _find_summary_paragraph(blocks: list[dict]) -> list[dict] | None:
    """Groups consecutive, vertically-adjacent, same-column blocks into runs
    and returns the first one that reads as actual prose -- deliberately
    NOT "the one block over N characters" on its own, because block
    granularity isn't consistent across CV generators: on some files each
    visual LINE comes back as its own PyMuPDF block, so the paragraph only
    shows up once several get run together; on others (confirmed on a real
    Canva-style export) the whole paragraph is already one block with its
    lines pre-joined. Requiring 2+ blocks to accept a run handled the first
    case but silently discarded the second -- a single-block paragraph never
    flushed, so the scan fell through to the next thing that happened to
    look like multi-line prose (on that real file: the address/phone/email
    block under "contactez-moi"), and that's what got overwritten instead.
    Length alone is the right test in both cases.

    Starting a run needs a real prose-looking line (several words, not
    all-caps, so a short label like "PROFIL" can't kick one off), but once
    a run is going, a short line still EXTENDS it (a paragraph's last
    wrapped line is often just two or three words -- "aux détails." --
    excluding it here would leave it un-redacted and orphaned below the
    tailored text). Only an all-caps heading (a new section starting) or a
    real gap in position ends a run."""
    ordered = sorted(blocks, key=lambda b: (b["bbox"][0] > 250, b["bbox"][1]))
    run: list[dict] = []

    def flush() -> list[dict] | None:
        combined = " ".join(r["text"] for r in run)
        return run if len(combined) > 80 else None

    for b in ordered:
        text = b["text"]
        starts_prose = len(text.split()) >= 4 and not text.isupper()
        continues = bool(run) and not text.isupper() and \
            abs(b["bbox"][0] - run[-1]["bbox"][0]) <= 5 and b["bbox"][1] - run[-1]["bbox"][3] <= 20
        if continues:
            run.append(b)
        elif starts_prose:
            found = flush()
            if found:
                return found
            run = [b]
        else:
            found = flush()
            if found:
                return found
            run = []
    return flush()


def _tailor_pdf_text(doc: fitz.Document, profile: dict, title: str, description: str) -> None:
    """Two full passes across every edit, not one pass per edit: every
    add_redact_annot() has to happen before the single apply_redactions()
    call, and every font registration + insert_text() has to happen after
    it. Interleaving (redact this zone, insert its text, redact the NEXT
    zone, insert...) was tried and confirmed broken -- a later
    apply_redactions() rebuilds the page's whole content stream and
    corrupts text inserted by an earlier one, not just breaks its own font."""
    page = doc[0]
    blocks = _text_blocks(page)
    edits: list[dict] = []  # each: bbox, new_text, spans

    run = _find_summary_paragraph(blocks)
    if run:
        summary_original = " ".join(b["text"] for b in run)
        x0 = min(b["bbox"][0] for b in run)
        y0 = min(b["bbox"][1] for b in run)
        x1 = max(b["bbox"][2] for b in run)
        y1 = max(b["bbox"][3] for b in run)
        new_summary = _tailor_summary_text(summary_original, profile, title, description)
        if new_summary != summary_original:
            edits.append({
                "bbox": fitz.Rect(x0 - 1, y0 - 2, x1 + 1, y1 + 2),
                "text": new_summary, "spans": run[0]["spans"],
            })

    # skill lines: swap the text of each currently-visible skill line for a
    # better-matching real skill the candidate has but that isn't shown yet
    # -- never touching count/position/formatting, one line in, one line out.
    skill_blocks = _find_skill_section_blocks(blocks)
    visible_skills = [b["text"] for b in skill_blocks]
    hidden_picks = _pick_hidden_skills(visible_skills, profile.get("skills", []), title, description,
                                        max_picks=min(2, len(visible_skills)))
    for old_skill, new_skill in zip(visible_skills, hidden_picks):
        block = next(b for b in skill_blocks if b["text"] == old_skill)
        edits.append({"bbox": fitz.Rect(*block["bbox"]), "text": new_skill, "spans": block["spans"]})

    if not edits:
        return

    # phase 1: decide fonts/backgrounds and redact everything, from the
    # page's original state
    for e in edits:
        e["bg"] = _sample_background(page, e["bbox"])
        e["fontfile"], e["fontsize"], e["color"] = _pick_font(page, e["spans"], e["text"])
        page.add_redact_annot(e["bbox"], fill=tuple(c / 255 for c in e["bg"]))
    page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_PIXELS)

    # phase 2: register fonts and insert all the new text, only now
    for e in edits:
        _wrap_and_insert(page, e["bbox"], e["text"], e["fontfile"], e["fontsize"], e["color"])


_SKILL_HEADER_WORDS = ("competence", "competences", "skill", "skills", "aptitude", "aptitudes")


def _is_skill_header(text: str) -> bool:
    # CV section headers are routinely letter-spaced ("C O M P É T E N C E S"),
    # so the keyword check has to run on the whitespace-stripped text, not
    # the raw one.
    if not text.isupper():
        return False
    squashed = normalize_text(text).replace(" ", "")
    return any(word in squashed for word in _SKILL_HEADER_WORDS)


def _find_skill_section_blocks(blocks: list[dict]) -> list[dict]:
    """Anchored on an actual "skills"/"compétences"-type section header,
    not just "looks short like the others" -- an earlier, width-only
    heuristic here matched the candidate's own name, address, and dates as
    "skill lines" on a real CV (confirmed directly), which would have
    overwritten the wrong text entirely. Collects the short, single-line,
    non-header blocks in the same column right after that header, stopping
    at the next all-caps header or a real gap in position."""
    ordered = sorted(blocks, key=lambda b: (b["bbox"][0] > 250, b["bbox"][1]))
    header_idx = next((i for i, b in enumerate(ordered) if _is_skill_header(b["text"])), None)
    if header_idx is None:
        return []

    header = ordered[header_idx]
    result = []
    prev = header
    for b in ordered[header_idx + 1:]:
        if (b["bbox"][0] > 250) != (header["bbox"][0] > 250):
            break  # crossed into a different column
        if b["text"].isupper():
            break  # next section header
        if b["bbox"][1] - prev["bbox"][3] > 25:
            break  # real vertical gap -- likely a different block entirely
        if 2 <= len(b["text"]) <= 40 and "\n" not in b["text"]:
            result.append(b)
        prev = b
    return result


# ---------------------------------------------------------------------------
# docx branch
# ---------------------------------------------------------------------------

def _ensure_docx_template(source_path: Path) -> Path:
    """One-time setup per uploaded .docx: find the paragraph that looks like
    the summary (the longest short prose paragraph, same heuristic as the
    PDF branch) and replace its text with a {{ cv_summary }} Jinja tag,
    saved as a separate template file -- the original upload stays untouched
    for reference, docxtpl renders fresh copies from the template from then
    on. docxtpl (not raw python-docx) specifically because Word/Docs
    routinely split one visually uniform sentence across several `run`
    objects even without a formatting change; docxtpl's tag rendering
    handles that run-splitting itself instead of this project re-solving it."""
    from docx import Document as _Document

    template_path = PROFILE_DIR / "template.docx"
    if template_path.exists():
        return template_path

    doc = _Document(str(source_path))
    target = max(doc.paragraphs, key=lambda p: len(p.text) if len(p.text) > 80 else 0, default=None)
    if target is None or len(target.text) <= 80:
        raise ValueError("no summary-like paragraph found in the .docx CV")
    for run in target.runs:
        run.text = ""
    target.runs[0].text = "{{ cv_summary }}"
    doc.save(str(template_path))
    return template_path


def _tailor_docx(profile: dict, title: str, description: str, offer_id: int) -> Path:
    from docxtpl import DocxTemplate

    source_path = Path(profile["cv_source_path"])
    template_path = _ensure_docx_template(source_path)

    from docx import Document as _Document
    original_summary = ""
    for p in _Document(str(source_path)).paragraphs:
        if len(p.text) > 80:
            original_summary = p.text
            break

    new_summary = _tailor_summary_text(original_summary, profile, title, description)

    tpl = DocxTemplate(str(template_path))
    tpl.render({"cv_summary": new_summary})
    out_dir = OUTPUT_DIR / str(offer_id)
    out_dir.mkdir(parents=True, exist_ok=True)
    rendered_docx = out_dir / "CV.docx"
    tpl.save(str(rendered_docx))

    _convert_docx_to_pdf(rendered_docx, out_dir)
    rendered_docx.unlink(missing_ok=True)
    return out_dir / "CV.pdf"


def _convert_docx_to_pdf(docx_path: Path, out_dir: Path) -> None:
    """LibreOffice headless -- a real, new system dependency, not something
    pip installs (see README). A unique --env:UserInstallation profile per
    call avoids a known LibreOffice failure mode: concurrent headless
    invocations sharing the default profile can deadlock, and
    draft_letters_node can process several offers (each spawning its own
    conversion) in a single run."""
    profile_dir = Path(tempfile.gettempdir()) / f"hobot-soffice-{uuid.uuid4().hex}"
    subprocess.run(
        ["soffice", "--headless", f"-env:UserInstallation=file://{profile_dir}",
         "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
        check=True, capture_output=True, timeout=60,
    )


# ---------------------------------------------------------------------------
# pdf_image branch (no text layer at all)
# ---------------------------------------------------------------------------

LOCATE_SUMMARY_PROMPT = """This is a CV page image. Find the profile/summary paragraph only (the
short introductory paragraph near the top, not the work experience, education, or contact
details).

Reply with ONLY JSON: {{"found": true/false, "text": "verbatim transcription of that paragraph",
"x0": 0.0, "y0": 0.0, "x1": 1.0, "y1": 1.0}}

x0/y0/x1/y1 are the paragraph's bounding box as FRACTIONS of the page width/height (0.0 to 1.0),
not pixels."""


def _locate_summary_zone(page) -> dict | None:
    pix = page.get_pixmap(dpi=150)
    result = chat_json(LOCATE_SUMMARY_PROMPT, images=[pix.tobytes("png")])
    if not result.get("found"):
        return None
    w, h = page.rect.width, page.rect.height
    bbox = fitz.Rect(result["x0"] * w, result["y0"] * h, result["x1"] * w, result["y1"] * h)
    # cheap self-correction: re-render just the claimed crop and ask for a
    # yes/no confirmation before trusting a small local model's bbox guess
    crop = page.get_pixmap(dpi=150, clip=bbox)
    confirm = chat_json(
        'Does this cropped image show a CV\'s profile/summary paragraph? Reply with ONLY JSON: {"yes": true/false}',
        images=[crop.tobytes("png")],
    )
    if not confirm.get("yes"):
        return None
    return {"bbox": bbox, "text": result.get("text", "")}


def _tailor_pdf_image(doc: fitz.Document, profile: dict, title: str, description: str) -> bool:
    """Returns True if a real, ATS-readable edit was made. Deliberately does
    NOT attempt to reconstruct the whole CV as an editable template (a much
    harder, much less scoped task than anything else in this project, and a
    regenerated template is by definition not the same file) and does NOT
    just repaint pixels (an ATS parser gets nothing from pixels -- inserting
    a real PDF text object is the part that actually makes this branch
    meaningful). Font fidelity is not attempted here; this branch already
    can't promise visual fidelity and says so in the caller's notification."""
    page = doc[0]
    zone = _locate_summary_zone(page)
    if zone is None:
        return False

    bbox = zone["bbox"]
    new_summary = _tailor_summary_text(zone["text"], profile, title, description)

    bg = _sample_background(page, bbox)
    page.draw_rect(bbox, color=None, fill=tuple(c / 255 for c in bg))
    _wrap_and_insert(page, bbox, new_summary, FONTS_DIR / "Montserrat-Regular.ttf", 10.0, (0, 0, 0))
    return True


# ---------------------------------------------------------------------------
# public entry point
# ---------------------------------------------------------------------------

def cv_path(offer_id: int) -> Path:
    return OUTPUT_DIR / str(offer_id) / "CV.pdf"


def tailor_cv(offer_id: int, title: str, description: str) -> Path | None:
    """Shared entry point for both the automatic pipeline
    (graphs/discovery_graph.py::draft_letters_node) and the on-demand chat
    tool (graphs/chat_agent.py::adapter_cv). Never raises past this
    boundary -- a failed tailoring attempt must never break letter
    drafting or the scoring pipeline around it, same convention as
    tools/web_search.py::search_company returning "" on failure."""
    try:
        with get_connection() as conn:
            profile = _full_profile(conn)
        if not profile.get("cv_source_path"):
            return None

        fmt = profile["cv_format"]
        out_dir = OUTPUT_DIR / str(offer_id)
        out_dir.mkdir(parents=True, exist_ok=True)

        if fmt == "docx":
            return _tailor_docx(profile, title, description, offer_id)

        doc = fitz.open(profile["cv_source_path"])
        try:
            if fmt == "pdf_text":
                _tailor_pdf_text(doc, profile, title, description)
            elif fmt == "pdf_image":
                if not _tailor_pdf_image(doc, profile, title, description):
                    return None
            else:
                return None
            out_path = cv_path(offer_id)
            doc.save(str(out_path))
            return out_path
        finally:
            doc.close()
    except Exception:
        return None
